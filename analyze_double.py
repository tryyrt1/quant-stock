#!/usr/bin/env python3
"""
翻倍股启动前指标规律分析 v2
- 腾讯API获取K线，东方财富API获取板块/换手率/流通股本
"""
import requests, json, time, re, os, sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import Counter, defaultdict
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    os.system('pip install python-docx -q')
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

HEADERS = {"User-Agent": "Mozilla/5.0"}
OUTPUT_DIR = "/home/ubuntu/quant-stock"


def _f(v):
    try: return float(v)
    except: return 0.0

def _int(v):
    try: return int(float(v))
    except: return 0


# ===================== DATA FETCHING =====================

def get_sectors():
    """获取所有行业板块"""
    url = "http://push2.eastmoney.com/api/qt/clist/get"
    params = {"pn": 1, "pz": 200, "fs": "m:90+t:2", "fields": "f12,f14,f20", "fid": "f20"}
    try:
        r = requests.get(url, params=params, timeout=15, headers=HEADERS)
        items = r.json().get("data", {}).get("diff", {})
        if isinstance(items, dict): items = list(items.values())
        if not isinstance(items, list): items = []
        sectors = []
        for i in items:
            if not isinstance(i, dict): continue
            name = str(i.get("f14", ""))
            if not name or "ST" in name: continue
            sectors.append({"code": i["f12"], "name": name, "count": i.get("f20", 0)})
        return sectors[:120]
    except Exception as e:
        print(f"[ERROR] get_sectors: {e}")
        return []


def get_sector_stocks(sector_code, max_stocks=15):
    """获取板块内成分股"""
    url = "http://push2.eastmoney.com/api/qt/clist/get"
    params = {
        "pn": 1, "pz": max_stocks,
        "fs": f"b:{sector_code}+f:!50",
        "fields": "f12,f14,f2,f3,f5,f6,f8,f9,f20,f21",
        "fid": "f3", "po": 1
    }
    try:
        r = requests.get(url, params=params, timeout=15, headers=HEADERS)
        items = r.json().get("data", {}).get("diff", {})
        if isinstance(items, dict): items = list(items.values())
        if not isinstance(items, list): items = []
        stocks = []
        for i in items:
            if not isinstance(i, dict): continue
            code = str(i.get("f12", ""))
            name = str(i.get("f14", ""))
            if code.startswith("3") or code.startswith("688"): continue
            if "ST" in name.upper() or "*" in name: continue
            stocks.append({
                "code": code, "name": name,
                "price": i.get("f2", 0) or 0,
                "change_pct": i.get("f3", 0) or 0,
                "volume": i.get("f5", 0) or 0,
                "amount": i.get("f6", 0) or 0,
                "turnover_rate": i.get("f8", 0) or 0,
                "pe": i.get("f9", 0) or 0,
                "market_cap": i.get("f20", 0) or 0,
                "float_cap": i.get("f21", 0) or 0,
            })
        return stocks
    except Exception as e:
        return []


def get_circulating_shares(code):
    """获取流通股本(股) 从腾讯行情"""
    market = "sh" if code.startswith("6") else "sz"
    try:
        url = f"https://qt.gtimg.cn/q={market}{code}"
        r = requests.get(url, timeout=8, headers=HEADERS)
        r.encoding = "gbk"
        m = re.search(r'="(.+)"', r.text)
        if not m: return 0
        fields = m.group(1).split("~")
        # field 44 = 流通股本(股), field 45 = 总股本(股)
        if len(fields) > 45:
            return _int(fields[44])
        return 0
    except:
        return 0


def fetch_kline_tx(code, days=300):
    """从腾讯获取K线数据"""
    market = "sh" if code.startswith("6") else "sz"
    full = market + code
    url = f"https://web.ifzq.gtimg.cn/appstock/app/fqkline/get?param={full},day,,,{days},qfq"
    try:
        r = requests.get(url, timeout=10, headers=HEADERS)
        raw = r.json()
        data = raw.get("data", {})
        klines = data.get(full, {}).get("qfqday") or data.get(full, {}).get("day") or []
        result = []
        for k in klines:
            result.append({
                "date": k[0],
                "open": _f(k[1]), "close": _f(k[2]),
                "high": _f(k[3]), "low": _f(k[4]),
                "volume": _int(k[5]),  # 手
            })
        return result
    except Exception as e:
        return []


def calc_turnover(volume_shou, circulating_shares):
    """计算换手率(%)"""
    if circulating_shares <= 0:
        return 0
    return volume_shou * 100 / circulating_shares * 100


# ===================== ANALYSIS =====================

def analyze_stock(klines, circulating_shares):
    """
    分析单只股票
    返回 (is_doubled, analysis_dict)
    """
    if len(klines) < 120:
        return False, None

    closes = [k["close"] for k in klines]
    n = len(closes)

    # 计算均线
    ma20 = []
    for i in range(n):
        if i < 19: ma20.append(None)
        else: ma20.append(sum(closes[i-19:i+1]) / 20)

    ma60 = []
    for i in range(n):
        if i < 59: ma60.append(None)
        else: ma60.append(sum(closes[i-59:i+1]) / 60)

    # 近期最高 vs 历史最低（去掉最近30天）
    recent_max = max(closes[-60:])
    hist_min = min(closes[:n-30])

    # 必须翻倍
    if recent_max < hist_min * 1.8:
        return False, None

    # 找最低点位置
    min_idx = closes.index(hist_min) if hist_min in closes else 0

    # 从最低点后找突破点：价格站上MA20且趋势确立
    breakout_idx = None
    for i in range(min_idx + 5, n - 20):
        if ma20[i] is None or ma60[i] is None: continue
        if closes[i] > ma20[i] and ma20[i] > ma20[i-1]:
            future = closes[i:i+20]
            if len(future) < 15: continue
            if sum(1 for x in future if x > ma20[i]) >= len(future) * 0.6:
                breakout_idx = i
                break

    if breakout_idx is None:
        return False, None

    # 涨幅验证
    peak = max(closes[breakout_idx:])
    gain = (peak - closes[breakout_idx]) / closes[breakout_idx] * 100
    if gain < 80:
        return False, None

    # ===== 启动前分析区间: 突破前40天 =====
    pre_start = max(0, breakout_idx - 40)
    pre_end = breakout_idx
    pre = klines[pre_start:pre_end]

    if len(pre) < 10:
        return False, None

    pre_closes = [k["close"] for k in pre]
    pre_volumes = [k["volume"] for k in pre]

    # 计算换手率 (通过流通股本)
    pre_turnovers = []
    for k in pre:
        tr = calc_turnover(k["volume"], circulating_shares)
        if tr > 0:
            pre_turnovers.append(tr)

    # 连阳统计
    max_consec_up = 0; cur = 0
    for k in pre:
        if k["close"] >= k["open"]: cur += 1; max_consec_up = max(max_consec_up, cur)
        else: cur = 0

    last_consec_up = 0
    for k in reversed(pre):
        if k["close"] >= k["open"]: last_consec_up += 1
        else: break

    max_consec_down = 0; cur = 0
    for k in pre:
        if k["close"] < k["open"]: cur += 1; max_consec_down = max(max_consec_down, cur)
        else: cur = 0

    # 换手率趋势
    avg_turn = sum(pre_turnovers) / len(pre_turnovers) if pre_turnovers else 0
    mid = len(pre_turnovers) // 2
    first_avg = sum(pre_turnovers[:mid]) / mid if mid > 0 else 0
    second_avg = sum(pre_turnovers[mid:]) / (len(pre_turnovers) - mid) if (len(pre_turnovers) - mid) > 0 else 0
    turn_trend = "递增" if second_avg > first_avg * 1.15 else ("递减" if second_avg < first_avg * 0.85 else "平稳")

    # 阳线比例
    changes = [k["close"] - k["open"] for k in pre]
    up_days = sum(1 for c in changes if c > 0)
    down_days = sum(1 for c in changes if c < 0)
    up_ratio = up_days / (up_days + down_days) * 100 if (up_days + down_days) > 0 else 50

    # 大阳线(涨>3%)
    pct_changes = [(k["close"] - k["open"]) / k["open"] * 100 for k in pre]
    big_up = sum(1 for c in pct_changes if c > 3)
    big_up_ratio = big_up / len(pct_changes) * 100 if pct_changes else 0

    # 均涨幅
    avg_pct = sum(pct_changes) / len(pct_changes) if pct_changes else 0

    # 最大回撤
    max_dd = 0; peak_p = pre[0]["close"]
    for k in pre:
        peak_p = max(peak_p, k["close"])
        dd = (peak_p - k["close"]) / peak_p * 100
        max_dd = max(max_dd, dd)

    # 成交量倍率 (启动前 vs 更早期)
    earlier = klines[max(0, pre_start-30):pre_start]
    earlier_vol = sum(k["volume"] for k in earlier) / len(earlier) if earlier else 1
    pre_vol_avg = sum(pre_volumes) / len(pre_volumes) if pre_volumes else 1
    vol_ratio = pre_vol_avg / earlier_vol if earlier_vol > 0 else 1

    # 尾端放量
    last10 = pre[-10:] if len(pre) >= 10 else pre
    before = pre[:-10] if len(pre) >= 10 else []
    last10_vol = sum(k["volume"] for k in last10) / len(last10) if last10 else 1
    before_vol = sum(k["volume"] for k in before) / len(before) if before else 1
    last_vol_ratio = last10_vol / before_vol if before_vol > 0 else 1

    # 价格相对MA位置
    last_close = pre[-1]["close"]
    ema20_n = sum(pre_closes[-20:]) / 20 if len(pre_closes) >= 20 else sum(pre_closes) / len(pre_closes)
    ma60_pre = [x for x in ma60[pre_start:pre_end] if x is not None]
    ema60_n = sum(ma60_pre) / len(ma60_pre) if ma60_pre else last_close

    return True, {
        "breakout_idx": breakout_idx,
        "breakout_date": klines[breakout_idx]["date"],
        "breakout_price": round(closes[breakout_idx], 2),
        "peak_price": round(peak, 2),
        "total_gain": round(gain, 0),
        "min_price": round(hist_min, 2),
        "min_date": klines[min_idx]["date"],
        "days_pre": len(pre),
        "avg_turnover": round(avg_turn, 2),
        "turnover_trend": turn_trend,
        "max_consec_up": max_consec_up,
        "last_consec_up": last_consec_up,
        "max_consec_down": max_consec_down,
        "up_ratio": round(up_ratio, 1),
        "avg_change": round(avg_pct, 2),
        "big_up_ratio": round(big_up_ratio, 1),
        "max_drawdown": round(max_dd, 1),
        "vol_ratio": round(vol_ratio, 2),
        "last_vol_ratio": round(last_vol_ratio, 2),
        "price_to_ma20_pre": round((last_close - ema20_n) / ema20_n * 100, 1),
        "pre_start_date": klines[pre_start]["date"],
        "pre_end_date": klines[pre_end-1]["date"],
    }


# ===================== REPORT GEN =====================

def gen_report(doubled_stocks, sectors, total_processed):
    """生成Word报告"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)

    total = len(doubled_stocks)
    if total == 0:
        doc.add_heading("分析结果", level=1)
        doc.add_paragraph("未发现符合条件的翻倍股。")
        doc.save(os.path.join(OUTPUT_DIR, "翻倍股分析报告.docx"))
        return

    pre_metrics = [s["pre"] for s in doubled_stocks]

    # 统计
    stats = {
        "total": total,
        "avg_turnover": sum(m["avg_turnover"] for m in pre_metrics) / total,
        "max_consec_up": sum(m["max_consec_up"] for m in pre_metrics) / total,
        "last_consec_up": sum(m["last_consec_up"] for m in pre_metrics) / total,
        "consec_down": sum(m["max_consec_down"] for m in pre_metrics) / total,
        "up_ratio": sum(m["up_ratio"] for m in pre_metrics) / total,
        "big_up_ratio": sum(m["big_up_ratio"] for m in pre_metrics) / total,
        "avg_change": sum(m["avg_change"] for m in pre_metrics) / total,
        "max_dd": sum(m["max_drawdown"] for m in pre_metrics) / total,
        "vol_ratio": sum(m["vol_ratio"] for m in pre_metrics) / total,
        "last_vol_ratio": sum(m["last_vol_ratio"] for m in pre_metrics) / total,
        "avg_gain": sum(s["breakout"]["total_gain"] for s in doubled_stocks) / total,
        "price_to_ma20": sum(m["price_to_ma20"] for m in pre_metrics) / total,
    }
    trend_dist = Counter(m["turnover_trend"] for m in pre_metrics)

    # === 封面 ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("翻倍股启动前指标规律分析报告")
    run.bold = True; run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(12); run.font.color.rgb = RGBColor(100,100,100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"覆盖板块: {len(sectors)}个 | 分析股票: {total_processed}只 | 发现翻倍股: {total}只")
    run.font.size = Pt(11); run.font.color.rgb = RGBColor(100,100,100)

    doc.add_page_break()

    # === 核心发现 ===
    doc.add_heading("一、核心发现", level=1)
    findings = [
        f"在{len(sectors)}个行业板块中，共分析{total_processed}只股票，发现{total}只翻倍股。",
        f"启动前日均换手率: {stats['avg_turnover']:.2f}%，"
        f"换手率趋势分布: 递增{trend_dist.get('递增',0)}只、平稳{trend_dist.get('平稳',0)}只、递减{trend_dist.get('递减',0)}只。",
        f"最长连阳均值: {stats['max_consec_up']:.1f}天，最后一段连阳: {stats['last_consec_up']:.1f}天。",
        f"阳线占比: {stats['up_ratio']:.1f}%（>50%表示多头占优）。",
        f"大阳线(涨>3%)占比: {stats['big_up_ratio']:.1f}%。",
        f"成交量倍率(相对更早期): {stats['vol_ratio']:.2f}倍，尾端放量倍率: {stats['last_vol_ratio']:.2f}倍。",
        f"最大回撤均值: {stats['max_dd']:.1f}%。",
        f"价格相对MA20: {stats['price_to_ma20']:.1f}%，启动后平均涨幅: {stats['avg_gain']:.0f}%。",
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # === 统计表 ===
    doc.add_heading("二、关键指标统计表", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(["指标", "均值", "说明"]):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.bold = True

    rows_data = [
        ("换手率(%)", f"{stats['avg_turnover']:.2f}", "启动前日均换手率"),
        ("最长连阳(天)", f"{stats['max_consec_up']:.1f}", "启动前最长连续阳线"),
        ("最后连阳(天)", f"{stats['last_consec_up']:.1f}", "启动前最后一段连阳"),
        ("最长连阴(天)", f"{stats['consec_down']:.1f}", "启动前最长连续阴线"),
        ("阳线比例(%)", f"{stats['up_ratio']:.1f}", "阳线>50%表示主力吸筹"),
        ("大阳线比例(%)", f"{stats['big_up_ratio']:.1f}", "涨幅>3%交易日占比"),
        ("日均涨跌幅(%)", f"{stats['avg_change']:.2f}", "启动前日均价格变动"),
        ("最大回撤(%)", f"{stats['max_dd']:.1f}", "启动前最大回调幅度"),
        ("成交量倍率", f"{stats['vol_ratio']:.2f}", "启动前均量/更早期均量"),
        ("尾端放量倍率", f"{stats['last_vol_ratio']:.2f}", "最后10天/之前均量"),
        ("价格距MA20(%)", f"{stats['price_to_ma20']:.1f}", "启动前相对20日线"),
        ("启动后涨幅(%)", f"{stats['avg_gain']:.0f}", "从启动点到近期最高"),
    ]
    for name, val, desc in rows_data:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = val
        row.cells[2].text = desc
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === 翻倍股详情 ===
    doc.add_heading("三、翻倍股详情", level=1)
    doubled_stocks.sort(key=lambda x: x["breakout"]["total_gain"], reverse=True)

    for i, s in enumerate(doubled_stocks):
        if i >= 50:
            doc.add_paragraph(f"（共{total}只，前50只已展示，余{total-50}只略）")
            break
        bp = s["breakout"]; pre = s["pre"]
        doc.add_heading(f"{i+1}. {s['name']}({s['code']}) — {s['sector']}", level=2)
        doc.add_paragraph(
            f"涨幅: {bp['total_gain']:.0f}% | 启动日: {bp['breakout_date']} | "
            f"启动价: {bp['breakout_price']} | 最高: {bp['peak_price']} | "
            f"最低: {bp['min_price']}({bp['min_date']})"
        )
        doc.add_paragraph(
            f"启动前{pre['days_analyzed']}天分析: "
            f"换手率{pre['avg_turnover']}%({pre['turnover_trend']}) | "
            f"连阳{pre['max_consec_up']}天/最后{pre['last_consec_up']}天 | "
            f"阳线{pre['up_ratio']}% | 大阳线{pre['big_up_ratio']}%"
        )
        doc.add_paragraph(
            f"量比{pre['vol_ratio']} | 尾端放量{pre['last_vol_ratio']} | "
            f"回撤{pre['max_drawdown']}% | 距MA20:{pre['price_to_ma20']}%"
        )

    doc.add_page_break()

    # === 规律总结 ===
    doc.add_heading("四、规律总结与操作建议", level=1)

    conclusions = []

    turn = stats['avg_turnover']
    if turn > 3:
        conclusions.append("1. 换手率特征: 启动前换手率较高(>3%)，主力建仓明显，筹码交换活跃。")
    elif turn > 1:
        conclusions.append("1. 换手率特征: 启动前换手率温和(1-3%)，主力悄悄吸筹，未引起市场关注。")
    else:
        conclusions.append("1. 换手率特征: 启动前换手率偏低(<1%)，主力高度控盘。")

    if trend_dist.get('递增', 0) > trend_dist.get('递减', 0):
        conclusions.append("2. 换手率趋势: 多数呈递增，越临近启动越活跃，典型的资金进场信号。")
    else:
        conclusions.append("2. 换手率趋势: 换手率平稳或递减，主力控盘度高。")

    conclusions.append(f"3. 连阳特征: 最长连阳均值{stats['max_consec_up']:.1f}天，最后一段{stats['last_consec_up']:.1f}天。")
    conclusions.append(f"4. 多空力量: 阳线比例{stats['up_ratio']:.1f}%，多头在启动前已占优。")
    conclusions.append(f"5. 放量信号: 启动前成交量是更早期的{stats['vol_ratio']:.2f}倍，尾端放量{stats['last_vol_ratio']:.2f}倍。")
    conclusions.append(f"6. 洗盘特征: 启动前最大回撤{stats['max_dd']:.1f}%，主力有洗盘动作。")
    conclusions.append(f"7. 价格位置: 启动前在MA20下方{abs(stats['price_to_ma20']):.1f}%，处于相对低位。")

    for c in conclusions:
        doc.add_paragraph(c)

    doc.add_paragraph("")

    doc.add_heading("操作关注信号", level=2)
    tips = [
        "换手率温和放大且递增",
        "出现5天以上连阳",
        "阳线比例>55%",
        "成交量相对前期放大1.5倍以上",
        "尾端进一步放量",
        "价格在MA20附近但未远离",
        "多个信号共振时确定性更高",
    ]
    for t in tips:
        doc.add_paragraph(t, style='List Bullet')

    # 保存
    out_path = os.path.join(OUTPUT_DIR, "翻倍股分析报告.docx")
    doc.save(out_path)
    return out_path


# ===================== MAIN =====================

def main():
    print("=" * 60)
    print("翻倍股启动前指标规律分析")
    print(f"开始: {datetime.now().strftime('%H:%M:%S')}")
    print("=" * 60)

    # 1. 获取板块
    print("\n[1] 获取板块列表...")
    sectors = get_sectors()
    print(f"  获取到 {len(sectors)} 个板块")
    if not sectors:
        print("  ERROR: 无板块数据"); return

    # 2. 遍历板块获取成分股
    print("\n[2] 遍历板块获取股票...")
    all_candidates = []
    for i, sec in enumerate(sectors):
        stocks = get_sector_stocks(sec["code"])
        for s in stocks:
            all_candidates.append({"sector": sec["name"], **s})
        if (i+1) % 20 == 0:
            print(f"  {i+1}/{len(sectors)}板块, 累计{len(all_candidates)}只")
    print(f"  共 {len(all_candidates)} 只候选股")

    # 去重
    seen = set()
    unique = []
    for c in all_candidates:
        if c["code"] not in seen:
            seen.add(c["code"])
            unique.append(c)
    print(f"  去重后 {len(unique)} 只")

    # 3. 获取K线 + 分析
    print(f"\n[3] 获取K线并分析翻倍股...")

    doubled = []
    failed_kline = 0
    failed_analyze = 0

    for i, c in enumerate(unique):
        code = c["code"]
        klines = fetch_kline_tx(code, 300)
        if len(klines) < 100:
            failed_kline += 1
            continue

        # 获取流通股本（用于换手率计算）
        circ_shares = get_circulating_shares(code)
        if circ_shares <= 0:
            circ_shares = c.get("float_cap", 0) * 10000  # 东方财富返回的是万元

        ok, analysis = analyze_stock(klines, circ_shares)
        if not ok:
            failed_analyze += 1
            continue

        analysis["code"] = code
        analysis["name"] = c["name"]
        analysis["sector"] = c["sector"]
        analysis["current_price"] = c["price"]
        analysis["current_turnover"] = c["turnover_rate"]

        # 重组为嵌套结构
        doubled.append({
            "name": c["name"],
            "code": code,
            "sector": c["sector"],
            "breakout": {
                "breakout_date": analysis.get("breakout_date"),
                "breakout_price": analysis.get("breakout_price"),
                "peak_price": analysis.get("peak_price"),
                "total_gain": analysis.get("total_gain"),
                "min_price": analysis.get("min_price"),
                "min_date": analysis.get("min_date"),
            },
            "pre": {
                "days_analyzed": analysis.get("days_pre"),
                "avg_turnover": analysis.get("avg_turnover"),
                "turnover_trend": analysis.get("turnover_trend"),
                "max_consec_up": analysis.get("max_consec_up"),
                "last_consec_up": analysis.get("last_consec_up"),
                "max_consec_down": analysis.get("max_consec_down"),
                "up_ratio": analysis.get("up_ratio"),
                "avg_change": analysis.get("avg_change"),
                "big_up_ratio": analysis.get("big_up_ratio"),
                "max_drawdown": analysis.get("max_drawdown"),
                "vol_ratio": analysis.get("vol_ratio"),
                "last_vol_ratio": analysis.get("last_vol_ratio"),
                "price_to_ma20": analysis.get("price_to_ma20_pre"),
                "pre_start_date": analysis.get("pre_start_date"),
                "pre_end_date": analysis.get("pre_end_date"),
            },
        })

        if len(doubled) <= 5:
            print(f"  ! 翻倍: {c['name']}({code}) {c['sector']} "
                  f"涨{analysis['total_gain']:.0f}% 启动{analysis['breakout_date']}")

        if (i+1) % 50 == 0:
            print(f"  进度 {i+1}/{len(unique)} 翻倍{len(doubled)}只")

    print(f"\n  完成: kline失败{failed_kline} 分析未通过{failed_analyze}")
    print(f"  翻倍股: {len(doubled)} 只")

    if not doubled:
        print("  未发现翻倍股")
        # 仍然生成空报告
        gen_report([], sectors, len(unique))
        return

    # 4. 生成报告
    print(f"\n[4] 生成Word报告...")
    path = gen_report(doubled, sectors, len(unique))
    print(f"  报告: {path}")

    # 汇总
    print("\n" + "=" * 60)
    print(f"板块: {len(sectors)} | 分析: {len(unique)}只 | 翻倍股: {len(doubled)}只")
    print(f"报告: {path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
